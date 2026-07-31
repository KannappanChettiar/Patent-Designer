"""Builds the submit-ready .docx from a numbered PatentDocument using
python-docx. All formatting is applied through real Word styles defined
once in `_setup_styles`, never as ad-hoc per-run formatting scattered
through the document.
"""

import os
import re

from PIL import Image as PILImage
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from models import PatentDocument, Subsection
from normalizer import EQN_MARK, ITALIC_MARK

TAG_RE = re.compile(r"^(\[\d{4}\])\s?(.*)$", re.S)
BULLET_PREFIX = "\u2022 "
NUMBERED_ITEM_RE = re.compile(r"^\d+\.\s")


# ---------------------------------------------------------------------------
# Styles & page setup
# ---------------------------------------------------------------------------

def _get_or_add_style(document, name, kind=WD_STYLE_TYPE.PARAGRAPH):
    styles = document.styles
    try:
        return styles[name]
    except KeyError:
        return styles.add_style(name, kind)


def _clear_style_paragraph_borders(style):
    ppr = style.element.find(qn("w:pPr"))
    if ppr is None:
        return
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is not None:
        ppr.remove(pbdr)


def _force_times_new_roman(style, size_pt: int, bold: bool = False, italic: bool = False):
    style.font.name = "Times New Roman"
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.italic = italic
    style.font.underline = False
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), "Times New Roman")


def _style_run(run, size_pt: int, bold: bool = False, italic: bool = False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)


def _setup_styles(document):
    styles = document.styles

    normal = styles["Normal"]
    _force_times_new_roman(normal, 10)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Times New Roman")

    title = _get_or_add_style(document, "Title")
    title.base_style = normal
    _force_times_new_roman(title, 14, bold=True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(24)
    _clear_style_paragraph_borders(title)

    h1 = styles["Heading 1"]
    h1.base_style = normal
    _force_times_new_roman(h1, 10, bold=True)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.keep_together = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)

    h2 = styles["Heading 2"]
    h2.base_style = normal
    _force_times_new_roman(h2, 10, bold=True)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.keep_together = True
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(6)

    caption = _get_or_add_style(document, "Caption")
    caption.base_style = normal
    _force_times_new_roman(caption, 10, italic=True)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _setup_page(document):
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)


def _add_page_number_footer(document):
    section = document.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)

    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _fit_image_dimensions(image_path: str, max_width_cm: float = 14.0, max_height_cm: float = 20.0):
    try:
        with PILImage.open(image_path) as image:
            width_px, height_px = image.size
    except Exception:
        return Cm(max_width_cm), Cm(max_height_cm)

    if width_px <= 0 or height_px <= 0:
        return Cm(max_width_cm), Cm(max_height_cm)

    width_scale = max_width_cm / width_px
    height_scale = max_height_cm / height_px
    scale = min(width_scale, height_scale, 1.0)
    return Cm(width_px * scale), Cm(height_px * scale)


# ---------------------------------------------------------------------------
# Content helpers
# ---------------------------------------------------------------------------

def _add_runs_with_italics(paragraph, text):
    parts = text.split(ITALIC_MARK)
    # parts alternate: normal, italic, normal, italic, ...
    for i, part in enumerate(parts):
        if part == "":
            continue
        run = paragraph.add_run(part)
        _style_run(run, 10, italic=(i % 2 == 1))


def _add_heading1(document, text):
    paragraph = document.add_paragraph(style="Heading 1")
    run = paragraph.add_run(text.upper())
    _style_run(run, 10, bold=True)


def _split_tag(numbered_text):
    m = TAG_RE.match(numbered_text)
    if m:
        return m.group(1), m.group(2)
    return "", numbered_text


def _add_body_paragraph(document, numbered_text):
    tag, rest = _split_tag(numbered_text)

    if rest.startswith(EQN_MARK) and rest.endswith(EQN_MARK):
        inner = rest[len(EQN_MARK):-len(EQN_MARK)]
        p = document.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if tag:
            run = p.add_run(tag + "  ")
            _style_run(run, 10)
        _add_runs_with_italics(p, inner)
        return

    lines = rest.split("\n") if rest else [""]
    for i, line in enumerate(lines):
        if line == "":
            continue
        is_list_item = line.startswith(BULLET_PREFIX) or bool(NUMBERED_ITEM_RE.match(line))
        p = document.add_paragraph(style="Normal")
        if is_list_item:
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if i == 0 and tag:
            run = p.add_run(tag + " ")
            _style_run(run, 10)
        _add_runs_with_italics(p, line)


def _add_subsection(document, sub: Subsection):
    if sub.title.strip():
        paragraph = document.add_paragraph(style="Heading 2")
        run = paragraph.add_run(sub.title)
        _style_run(run, 10, bold=True)
    if not sub.paragraphs:
        return
    for para in sub.paragraphs:
        _add_body_paragraph(document, para)


# ---------------------------------------------------------------------------
# Main build
# ---------------------------------------------------------------------------

def build_document(doc: PatentDocument, output_path: str) -> str:
    document = Document()
    _setup_styles(document)
    _setup_page(document)
    _add_page_number_footer(document)

    title_paragraph = document.add_paragraph(style="Title")
    title_run = title_paragraph.add_run(doc.title.upper())
    _style_run(title_run, 14, bold=True)

    _add_heading1(document, "Field of the Invention")
    for para in doc.field_of_invention:
        _add_body_paragraph(document, para)

    _add_heading1(document, "Background of the Invention")
    for sub in doc.background:
        _add_subsection(document, sub)

    _add_heading1(document, "Summary of the Invention")
    for para in doc.summary:
        _add_body_paragraph(document, para)

    _add_heading1(document, "Brief Description of the Drawings")
    for fig in doc.brief_description_of_drawings:
        p = document.add_paragraph(style="Caption")
        run = p.add_run(f"FIG. {fig.number} \u2014 {fig.caption}")
        _style_run(run, 10, italic=True)

    _add_heading1(document, "Detailed Description of the Invention")
    for sub in doc.detailed_description:
        _add_subsection(document, sub)

    document.add_page_break()
    _add_heading1(document, "Claims")
    for claim in doc.claims:
        p = document.add_paragraph(style="Normal")
        run = p.add_run(f"Claim {claim.number}. {claim.text}")
        _style_run(run, 10)

    document.add_page_break()
    _add_heading1(document, "Abstract")
    for para in doc.abstract.split("\n\n"):
        if para:
            _add_body_paragraph(document, para)

    document.add_page_break()
    _add_heading1(document, "Figures")
    for fig in doc.figures:
        if fig.image_path and os.path.exists(fig.image_path):
            width, height = _fit_image_dimensions(fig.image_path)
            document.add_picture(fig.image_path, width=width, height=height)
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save(output_path)
    return output_path
